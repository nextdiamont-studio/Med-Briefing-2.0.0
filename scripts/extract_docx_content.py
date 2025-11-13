# -*- coding: utf-8 -*-
"""
Script para extrair conteúdo de arquivos .docx da pasta Materiais
e preparar análise para otimização de prompts de IA
"""
import sys
import os
from pathlib import Path
import json

# Configurar encoding para UTF-8
if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8')

def extract_docx_content(file_path):
    """Extrai conteúdo de arquivo .docx usando python-docx"""
    try:
        import docx
    except ImportError:
        print("Modulo python-docx nao encontrado. Instalando...")
        import subprocess
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        import docx

    try:
        doc = docx.Document(str(file_path))

        content = {
            "filename": os.path.basename(str(file_path)),
            "paragraphs": [],
            "tables": []
        }

        # Extrair parágrafos
        for para in doc.paragraphs:
            if para.text.strip():
                content["paragraphs"].append(para.text.strip())

        # Extrair tabelas
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            content["tables"].append(table_data)

        return content
    except Exception as e:
        print(f"Erro ao processar {file_path}: {str(e)}")
        return None

def main():
    # Diretório de materiais
    materiais_dir = Path(r"C:\Users\fclpa\Downloads\Med Briefing 2.0\med-briefing\Materiais")

    # Encontrar todos os arquivos .docx
    docx_files = list(materiais_dir.glob("*.docx"))

    all_content = {}

    print(f"Encontrados {len(docx_files)} arquivos .docx para processar\n")

    for docx_file in docx_files:
        print(f"Processando: {docx_file.name}")
        content = extract_docx_content(docx_file)
        if content:
            all_content[docx_file.name] = content
            print(f"  [OK] {len(content['paragraphs'])} paragrafos extraidos")
            print(f"  [OK] {len(content['tables'])} tabelas extraidas")

    # Salvar conteúdo extraído em JSON
    output_file = materiais_dir.parent / "scripts" / "materiais_extracted.json"
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(all_content, f, ensure_ascii=False, indent=2)

    print(f"\n[SUCESSO] Conteudo extraido salvo em: {output_file}")

    # Criar resumo analítico
    print("\n" + "="*80)
    print("RESUMO ANALITICO DOS MATERIAIS")
    print("="*80)

    for filename, content in all_content.items():
        print(f"\n[ARQUIVO] {filename}")
        print(f"   Total de paragrafos: {len(content['paragraphs'])}")
        print(f"   Total de tabelas: {len(content['tables'])}")

        # Primeiras linhas para contexto
        if content['paragraphs']:
            print(f"   Inicio: {content['paragraphs'][0][:100]}...")

    return all_content

if __name__ == "__main__":
    main()
